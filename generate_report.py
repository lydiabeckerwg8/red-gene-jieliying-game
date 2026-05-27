from __future__ import annotations

import json
import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
INDEX = ROOT / "index.html"
README = ROOT / "README.md"
OUTPUT = ROOT / "红色基因接力营网页小游戏项目报告.docx"


def extract_missions(html: str) -> list[dict]:
    match = re.search(r"const\s+missions\s*=\s*(\[.*?\]);\s*\n\s*const\s+routeStations", html, re.S)
    if not match:
        raise RuntimeError("未能在 index.html 中找到 const missions 数据。")

    js = f"""
const missions = {match.group(1)};
console.log(JSON.stringify(missions));
"""
    result = subprocess.run(
        ["node", "-e", js],
        cwd=ROOT,
        text=True,
        encoding="utf-8",
        capture_output=True,
        check=True,
    )
    return json.loads(result.stdout)


def get_project_meta(html: str, readme: str) -> dict:
    def first(pattern: str, default: str = "") -> str:
        found = re.search(pattern, html, re.S)
        return re.sub(r"\s+", " ", found.group(1)).strip() if found else default

    page_url = ""
    url_match = re.search(r"https?://[^\s`]+", readme)
    if url_match:
        page_url = url_match.group(0)

    css_count = len(re.findall(r"^\s*[.#a-zA-Z][^{]+\{", html, re.M))
    function_count = len(re.findall(r"\bfunction\s+\w+\s*\(", html))

    return {
        "title": first(r"<title>(.*?)</title>", "红色基因接力营"),
        "lead": first(r'<p class="lead">(.*?)</p>'),
        "page_url": page_url,
        "line_count": len(html.splitlines()),
        "css_count": css_count,
        "function_count": function_count,
    }


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_paragraph(doc: Document, text: str, first_line: bool = True):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.5
    if first_line:
        paragraph.paragraph_format.first_line_indent = Pt(21)
    run = paragraph.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(11)
    return paragraph


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.line_spacing = 1.4
    run = paragraph.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)


def option_summary(mission: dict) -> str:
    if mission["type"] == "choice":
        return mission["options"][mission["answer"]]
    if mission["type"] == "multi":
        return "、".join(item["text"] for item in mission["options"] if item.get("good"))
    if mission["type"] == "timeline":
        return " → ".join(mission["items"][i]["text"] for i in range(len(mission["items"])))
    return ""


def build_report(meta: dict, missions: list[dict]) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("《红色基因接力营》网页小游戏项目报告")
    run.bold = True
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(125, 22, 27)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("基于项目文件夹内容整理")
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(12)

    add_heading(doc, "一、项目概述", 1)
    add_paragraph(
        doc,
        f"《{meta['title']}》是一个以红色基因传播和趣味思政学习为核心的单文件静态网页小游戏。项目入口为 index.html，不依赖后端服务或额外安装流程，用户可以直接双击本地文件预览，也可以部署到 GitHub Pages、Netlify、Vercel、Cloudflare Pages 等静态站点平台。",
    )
    add_paragraph(
        doc,
        "游戏以“地图闯关”的方式组织学习内容，将红色遗址、历史人物、精神谱系、价值辨析和研学设计等知识点转化为可交互任务。玩家沿研学路线推进，完成任务后点亮节点、收集精神火种和主题卡片，兼顾知识性、参与感和课堂展示效果。",
    )

    add_heading(doc, "二、项目文件结构", 1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["文件/目录", "作用", "说明"]
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True)
        set_cell_shading(table.rows[0].cells[i], "F2D8B8")
    rows = [
        ("index.html", "网页游戏主体", "包含 HTML 结构、CSS 样式、JavaScript 逻辑和 20 道关卡数据。"),
        ("README.md", "项目说明", "记录项目名称、在线链接、本地预览方式和 GitHub Pages 发布步骤。"),
        (".nojekyll", "静态部署辅助", "用于 GitHub Pages 部署时跳过 Jekyll 处理。"),
        (".github/workflows/deploy-pages.yml", "自动部署配置", "用于通过 GitHub Actions 发布到 GitHub Pages。"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)

    add_heading(doc, "三、主题定位与教育目标", 1)
    add_paragraph(
        doc,
        "本项目面向课堂展示、课程作业或红色文化主题学习场景，主题明确，内容围绕“红色基因”的历史来源、精神内涵和现实转化展开。游戏不是单纯记忆年份和人物，而是通过题目讲解引导学习者理解理想信念、人民立场、实事求是、艰苦奋斗、服务人民、责任担当等关键词。",
    )
    add_bullet(doc, "知识目标：了解红色遗址、历史人物、红色精神谱系和重要历史节点。")
    add_bullet(doc, "能力目标：通过单选、多选和排序任务训练辨析、归纳和时间线梳理能力。")
    add_bullet(doc, "价值目标：把红色基因与当代青年的学习态度、集体责任和实践行动联系起来。")

    add_heading(doc, "四、游戏玩法与交互设计", 1)
    add_paragraph(
        doc,
        "游戏共设置 20 个关卡，左侧为主题海报、红色图片、研学路线图和研学护照，右侧为任务简报、关卡说明、题目区域、反馈区和操作按钮。用户完成题目后会获得精神火种、主题卡片与讲解，系统同步更新百分制得分、已点亮节点、收集卡片和进度条。",
    )
    add_bullet(doc, "单选题：点击选项后立即判断正误，正确选项和错误选项会用不同颜色提示。")
    add_bullet(doc, "多选题：用户选择多个关键词后提交，系统根据正确项和错误项计算得分。")
    add_bullet(doc, "排序题：用户通过上下移动按钮调整历史事件或精神主题的先后顺序。")
    add_bullet(doc, "地图机制：后续关卡默认锁定，完成当前任务后才会解锁下一站。")
    add_bullet(doc, "收集机制：每题通关后获得对应主题卡片，研学护照会记录收集进度。")
    add_bullet(doc, "总结机制：完成最后一题后生成红色基因接力宣言，呈现整体得分、等级评价和卡片成果。")

    add_heading(doc, "五、关卡内容梳理", 1)
    mission_table = doc.add_table(rows=1, cols=6)
    mission_table.style = "Table Grid"
    headers = ["序号", "类别", "标题", "题型", "核心问题", "答案/要点"]
    for i, header in enumerate(headers):
        set_cell_text(mission_table.rows[0].cells[i], header, bold=True)
        set_cell_shading(mission_table.rows[0].cells[i], "EBCB93")
    type_names = {"choice": "单选", "multi": "多选", "timeline": "排序"}
    for idx, mission in enumerate(missions, 1):
        cells = mission_table.add_row().cells
        values = [
            str(idx),
            mission.get("category", ""),
            mission.get("title", ""),
            type_names.get(mission.get("type"), mission.get("type", "")),
            mission.get("q", ""),
            option_summary(mission),
        ]
        for i, value in enumerate(values):
            set_cell_text(cells[i], value)

    add_heading(doc, "六、技术实现分析", 1)
    add_paragraph(
        doc,
        f"项目采用原生 HTML、CSS 和 JavaScript 实现，所有代码集中在 index.html 中，当前文件约 {meta['line_count']} 行。CSS 使用变量定义红色、金色、纸张色等主题色，并通过响应式布局适配桌面端和移动端。JavaScript 中包含约 {meta['function_count']} 个主要函数，负责关卡渲染、答题判断、精神火种统计、路线解锁、卡片收集、图片加载和最终总结。",
    )
    add_bullet(doc, "数据结构：missions 数组集中保存关卡类别、标题、导语、题目、选项、答案、讲解和题型。")
    add_bullet(doc, "状态管理：通过 index、rawScore、selected、answered、done、unlockedPrizes、answers 等变量记录当前关卡、得分、已选项、路线点亮状态、卡片收集和答题历史。")
    add_bullet(doc, "渲染逻辑：renderMission 根据题型调用 renderChoice、renderMulti 或 renderTimeline，动态生成对应交互组件。")
    add_bullet(doc, "评价逻辑：award 函数统一处理得分、徽章、反馈和进度更新，保证用户操作后能立即看到学习结果。")
    add_bullet(doc, "图片资源：图片通过 Wikimedia Commons 的 Special:FilePath 链接加载，并设置 fallback 备选图片，提高在线展示稳定性。")

    add_heading(doc, "七、视觉风格与界面特点", 1)
    add_paragraph(
        doc,
        "界面整体采用红色、深红、金色和宣纸色作为主视觉，契合红色文化主题。页面结构为左侧海报区与右侧任务区，左侧新增路线地图和研学护照，右侧新增接力任务简报。圆角卡片、路线节点、收集卡片、奖励提示和按钮反馈等元素让游戏具有更完整的闯关体验。",
    )
    add_bullet(doc, "主题色鲜明：红色体现红色文化主题，金色用于强调和进度反馈。")
    add_bullet(doc, "信息层级清楚：标题、导语、问题、选项、反馈区域分区明确。")
    add_bullet(doc, "响应式适配：在窄屏设备上自动切换为单列布局，便于手机浏览。")

    add_heading(doc, "八、运行与发布方式", 1)
    add_paragraph(
        doc,
        "本项目是纯静态网页，运行门槛较低。本地预览时直接双击 index.html 即可打开；在线发布时可以上传项目根目录到静态托管平台。README.md 中已经给出了 GitHub Pages 的发布步骤。",
    )
    if meta["page_url"]:
        add_paragraph(doc, f"项目当前记录的在线访问链接为：{meta['page_url']}", first_line=False)

    add_heading(doc, "九、项目亮点", 1)
    add_bullet(doc, "内容完整：20 道题覆盖红船精神、遵义会议、延安精神、井冈山道路、西柏坡精神、铁人精神、两弹一星精神等主题。")
    add_bullet(doc, "趣味增强：通过路线解锁、精神火种、主题卡片和研学护照，把普通答题转化为任务式闯关。")
    add_bullet(doc, "形式多样：单选、多选和排序题结合，避免答题形式单一。")
    add_bullet(doc, "反馈充分：答题后给出解析和奖励反馈，把知识点和现实启示连接起来。")
    add_bullet(doc, "部署方便：单文件项目便于提交、展示和线上发布。")

    add_heading(doc, "十、可优化方向", 1)
    add_bullet(doc, "可增加本地图片备份，降低外部图片链接失效或加载慢对展示效果的影响。")
    add_bullet(doc, "可增加开始页、成绩页导出、错题回顾等功能，让学习闭环更完整。")
    add_bullet(doc, "可把关卡数据拆分为 JSON 文件，便于后续维护和扩展更多主题。")
    add_bullet(doc, "可加入音效、动画或更丰富的角色化叙事，但应注意保持课堂展示的简洁性。")

    add_heading(doc, "十一、总结", 1)
    add_paragraph(
        doc,
        "总体来看，《红色基因接力营》把课程主题、网页技术和游戏化学习结合得比较紧密。它用轻量的前端实现承载了较完整的思政学习流程：先通过图片、路线和任务简报建立情境，再通过闯关题目完成知识辨析，随后用精神火种和卡片收集形成游戏反馈，最后通过接力宣言和总结评价强化理解。作为网页小游戏项目，它具备可运行、可展示、可部署和可继续扩展的特点，适合用于自然辩证法或相关课程中的数字化学习成果展示。",
    )

    doc.save(OUTPUT)


def main() -> None:
    html = INDEX.read_text(encoding="utf-8")
    readme = README.read_text(encoding="utf-8")
    missions = extract_missions(html)
    meta = get_project_meta(html, readme)
    build_report(meta, missions)
    print(OUTPUT)


if __name__ == "__main__":
    main()
